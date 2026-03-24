import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import os

# Crear carpetas
if not os.path.exists("cartas"):
    os.makedirs("cartas")

if not os.path.exists("graficas"):
    os.makedirs("graficas")

# Leer CSV
df = pd.read_csv("calificaciones.csv")

# Calcular promedio
df["Promedio"] = df.iloc[:, 2:7].mean(axis=1)

# 📊 GRÁFICA GENERAL
plt.figure()
plt.bar(df["Nombre"], df["Promedio"], color="pink")
plt.title("Promedio general de alumnos")
plt.xlabel("Alumnos")
plt.ylabel("Promedio")
plt.ylim(0, 10)
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("graficas/promedios_generales.png")
plt.close()

# Recomendación
recomendacion_larga = """Se ha identificado que el alumno presenta un desempeño académico por debajo del nivel esperado. 
Se recomienda establecer una rutina diaria de estudio en casa, reforzar los temas vistos en clase, así como fomentar hábitos 
de responsabilidad y organización. También es importante considerar el apoyo adicional mediante asesorías académicas o tutorías."""

# Recorrer alumnos
for _, alumno in df.iterrows():

    if alumno["Promedio"] < 6:

        nombre_archivo = alumno["Nombre"].replace(" ", "_")
        ruta_grafica = f"graficas/{nombre_archivo}.png"

        # 📊 GRÁFICA INDIVIDUAL
        materias = ["Matematicas", "Espanol", "Historia", "Ciencias", "Ingles"]
        calificaciones = [alumno[m] for m in materias]

        plt.figure()
        plt.bar(materias, calificaciones, color="pink")
        plt.ylim(0, 10)
        plt.title(f"{alumno['Nombre']}")
        plt.tight_layout()
        plt.savefig(ruta_grafica)
        plt.close()

        # 📄 DOCUMENTO
        doc = Document("plantilla_final_carta.docx")

        # 🔥 INSERTAR GRÁFICA ANTES DEL PROMEDIO
        for i, p in enumerate(doc.paragraphs):
            if "Promedio general" in p.text:

                # Insertar imagen después del texto
                img_parrafo = doc.paragraphs[i].insert_paragraph_before()
                run = img_parrafo.add_run()
                run.add_picture(ruta_grafica)

                break

        # 🔁 REEMPLAZOS (DESPUÉS)
        reemplazos = {
            "{{PADRE}}": alumno["Padre"],
            "{{NOMBRE}}": alumno["Nombre"],
            "{{MATEMATICAS}}": str(alumno["Matematicas"]),
            "{{ESPANOL}}": str(alumno["Espanol"]),
            "{{HISTORIA}}": str(alumno["Historia"]),
            "{{CIENCIAS}}": str(alumno["Ciencias"]),
            "{{INGLES}}": str(alumno["Ingles"]),
            "{{PROMEDIO}}": str(round(alumno["Promedio"], 2)),
            "{{RECOMENDACION}}": recomendacion_larga
        }

        # Párrafos
        for p in doc.paragraphs:
            for key, value in reemplazos.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        # Tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in reemplazos.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # Guardar
        doc.save(f"cartas/{nombre_archivo}.docx")

print("Cartas generadas correctamente.")